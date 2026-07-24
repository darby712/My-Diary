"""ReviewLens — Streamlit 버전.

앱 스토어 URL(최대 5개)을 넣으면 리뷰를 병렬 수집 → 전처리 → 분석 7종 → LLM 인사이트를 만들고,
CSV/Word/PDF로 내보낼 수 있는 단일 파일 앱. PRD.md의 F1~F6에 해당한다.

원래 FastAPI + Celery/Redis + React 구조(reviewlens/ 폴더)를 Streamlit 한 프로세스로 옮긴 버전이라
비동기 큐 대신 버튼 클릭 한 번에 동기로 전체 파이프라인을 실행하고, DB 대신 st.session_state에
결과를 담아둔다(브라우저 세션이 끝나면 사라짐 — 분석 이력을 남기려면 별도 저장소가 필요하다).
"""
import base64
import io
import json
import os
import re
from collections import Counter
from datetime import datetime
from urllib.parse import parse_qs, urlparse
import concurrent.futures

import pandas as pd
import streamlit as st

from google_play_scraper import app as gp_app, reviews as gp_reviews, Sort
from google_play_scraper.exceptions import NotFoundError
from konlpy.tag import Okt

# =====================================================================================
# 설정
# =====================================================================================

MAX_APPS = 5
DEFAULT_COUNT = 200
COLLECT_MAX_RETRIES = 3
COLLECT_RETRY_BACKOFF_SEC = 2
COLLECT_REQUEST_DELAY_SEC = 0.5
MONTHLY_NOISE_THRESHOLD = 5
TOP_N_FOR_COMMON = 100
WORDCLOUD_FONT_PATH = os.getenv("WORDCLOUD_FONT_PATH", "C:/Windows/Fonts/malgun.ttf")
REVIEW_LANG = os.getenv("REVIEW_LANG", "ko")
REVIEW_COUNTRY = os.getenv("REVIEW_COUNTRY", "kr")

STOPWORDS = {
    "앱", "것", "수", "좀", "너무", "정말", "진짜", "이거", "그냥", "거",
    "때", "더", "안", "못", "잘", "왜", "이게", "게", "도", "만", "은", "는",
    "이", "가", "을", "를", "의", "에", "로", "으로", "과", "와", "한", "하다",
}
POS_WORDS = ["좋아", "좋은", "편리", "편하", "만족", "간편", "최고", "감사", "빠르", "유용", "깔끔"]
NEG_WORDS = ["불편", "오류", "안돼", "안됨", "느리", "느려", "짜증", "최악", "버그", "먹통", "실패", "불만"]


def _get_secret(key: str, default: str = "") -> str:
    try:
        val = st.secrets.get(key)
        if val:
            return val
    except Exception:
        pass
    return os.getenv(key, default)


DEFAULT_ANTHROPIC_KEY = _get_secret("ANTHROPIC_API_KEY")
DEFAULT_ANTHROPIC_MODEL = _get_secret("ANTHROPIC_MODEL", "claude-opus-4-8")


# =====================================================================================
# F2: 리뷰 수집 (재시도·요청 딜레이·작성자명 마스킹 포함)
# =====================================================================================

_PACKAGE_ID_RE = re.compile(r"[a-zA-Z0-9_.]+")


def parse_package_id(url_or_id: str) -> str:
    """스토어 URL 또는 패키지명에서 패키지 id를 추출한다."""
    s = (url_or_id or "").strip()
    if urlparse(s).netloc == "play.google.com":
        q = parse_qs(urlparse(s).query)
        candidate = q.get("id", [None])[0]
        if candidate and _PACKAGE_ID_RE.fullmatch(candidate):
            return candidate
        raise ValueError(f"URL에서 유효한 패키지 id를 찾을 수 없습니다: {url_or_id}")
    if _PACKAGE_ID_RE.fullmatch(s):
        return s
    raise ValueError(f"패키지 id를 찾을 수 없습니다: {url_or_id}")


def mask_author(name) -> str | None:
    """작성자명 마스킹 (첫 글자만 남김) — 개인정보 최소화."""
    if not name:
        return None
    name = str(name)
    if len(name) <= 1:
        return "*"
    return name[0] + "*" * (len(name) - 1)


def fetch_app_info(package_id: str) -> dict:
    """앱 이름/아이콘 조회. 실패해도 패키지명으로 폴백."""
    try:
        info = gp_app(package_id, lang=REVIEW_LANG, country=REVIEW_COUNTRY)
        return {"name": info.get("title") or package_id, "icon_url": info.get("icon")}
    except Exception:
        return {"name": package_id, "icon_url": None}


def fetch_reviews(package_id: str, count: int) -> list[dict]:
    """리뷰를 수집해 정규화된 dict 리스트로 반환. 별점이 없거나 1~5 범위를 벗어나면 건너뛴다."""
    result, _ = gp_reviews(
        package_id, lang=REVIEW_LANG, country=REVIEW_COUNTRY, sort=Sort.NEWEST, count=count,
    )
    rows = []
    for r in result:
        content = (r.get("content") or "").strip()
        if not content:
            continue
        try:
            rating = int(r.get("score"))
        except (TypeError, ValueError):
            rating = None
        if rating is None or not (1 <= rating <= 5):
            continue
        rows.append({
            "rating": rating,
            "content": content,
            "author": mask_author(r.get("userName")),
            "created_at": r.get("at"),
        })
    return rows


def fetch_reviews_with_retry(package_id: str, count: int) -> list[dict]:
    """요청 제한·일시 오류에 대비해 재시도. 빈 결과(존재하지 않는 앱 등)도 명시적으로 실패시킨다."""
    import time as _time
    last_err: Exception | None = None
    for attempt in range(1, COLLECT_MAX_RETRIES + 1):
        try:
            rows = fetch_reviews(package_id, count)
            _time.sleep(COLLECT_REQUEST_DELAY_SEC)
            if not rows:
                raise ValueError(f"'{package_id}'에서 리뷰를 찾지 못했습니다 (존재하지 않는 앱이거나 리뷰가 없음).")
            return rows
        except NotFoundError:
            raise
        except Exception as e:  # noqa: BLE001 - 스토어 응답 오류는 재시도 대상으로 포괄 처리
            last_err = e
            if attempt < COLLECT_MAX_RETRIES:
                _time.sleep(COLLECT_RETRY_BACKOFF_SEC * attempt)
    raise last_err  # type: ignore[misc]


# =====================================================================================
# F3: 전처리 (중복/초단문 제거, 특수문자·이모지 제거, 명사 추출, 불용어 처리)
# =====================================================================================

@st.cache_resource(show_spinner="형태소 분석기(Okt) 준비 중... (처음 한 번만)")
def get_okt() -> Okt:
    return Okt()


def clean_text(text: str) -> str:
    text = re.sub(r"[^가-힣a-zA-Z0-9 ]", " ", str(text or ""))
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_nouns(clean_content: str) -> list[str]:
    nouns = get_okt().nouns(clean_content)
    return [w for w in nouns if len(w) > 1 and w not in STOPWORDS]


def preprocess_reviews(rows: list[dict]) -> list[dict]:
    seen: set[tuple] = set()
    out = []
    for row in rows:
        key = (row.get("app_id"), row.get("content"))
        if key in seen:
            continue
        seen.add(key)

        clean = clean_text(row.get("content"))
        if len(clean) <= 1:
            continue

        nouns = extract_nouns(clean)
        if not nouns:
            continue

        row = dict(row)
        row["clean_content"] = clean
        row["nouns"] = nouns
        out.append(row)
    return out


# =====================================================================================
# F4: 분석 7종
# =====================================================================================

def compute_sentiment(content: str) -> int:
    text = str(content or "")
    score = 0
    for w in POS_WORDS:
        if w in text:
            score += 1
    for w in NEG_WORDS:
        if w in text:
            score -= 1
    return score


def _word_counter(reviews: list[dict], app_id: int) -> Counter:
    c = Counter()
    for r in reviews:
        if r["app_id"] == app_id:
            c.update(r.get("nouns") or [])
    return c


def rating_distribution(apps: list[dict], reviews: list[dict]) -> dict:
    out = {}
    for a in apps:
        ratings = [r["rating"] for r in reviews if r["app_id"] == a["id"]]
        counts = {s: ratings.count(s) for s in range(1, 6)}
        avg = round(sum(ratings) / len(ratings), 2) if ratings else 0.0
        out[a["id"]] = {"counts": counts, "average": avg, "review_count": len(ratings)}
    return out


def wordcloud_frequencies(apps: list[dict], reviews: list[dict], top: int = 200) -> dict:
    return {a["id"]: dict(_word_counter(reviews, a["id"]).most_common(top)) for a in apps}


def top_keywords(apps: list[dict], reviews: list[dict], n: int = 20) -> dict:
    out = {}
    for a in apps:
        counter = _word_counter(reviews, a["id"])
        out[a["id"]] = [{"keyword": w, "count": c} for w, c in counter.most_common(n)]
    return out


def common_and_unique_keywords(apps: list[dict], reviews: list[dict], n: int = 10) -> dict:
    counters = {a["id"]: _word_counter(reviews, a["id"]) for a in apps}
    tops = {aid: set(w for w, _ in c.most_common(TOP_N_FOR_COMMON)) for aid, c in counters.items()}

    if len(apps) < 2:
        first = apps[0]["id"] if apps else None
        return {
            "common": [],
            "unique": {first: [w for w, _ in counters[first].most_common(n)]} if apps else {},
            "note": "앱이 1개뿐이라 공통 키워드를 계산할 수 없습니다.",
        }

    common_set = set.intersection(*tops.values())
    total = Counter()
    for c in counters.values():
        total.update(c)
    common_top = sorted(common_set, key=lambda w: total[w], reverse=True)[:n]

    unique = {}
    for a in apps:
        aid = a["id"]
        others = set()
        for other in apps:
            if other["id"] != aid:
                others |= tops[other["id"]]
        only = tops[aid] - others
        unique[aid] = sorted(only, key=lambda w: counters[aid][w], reverse=True)[:n]

    return {"common": common_top, "unique": unique, "note": None}


def positive_negative_keywords(apps: list[dict], reviews: list[dict], n: int = 10) -> dict:
    out = {}
    for a in apps:
        pos, neg = Counter(), Counter()
        for r in reviews:
            if r["app_id"] != a["id"]:
                continue
            if r["rating"] >= 4:
                pos.update(r.get("nouns") or [])
            elif r["rating"] <= 2:
                neg.update(r.get("nouns") or [])
        out[a["id"]] = {
            "positive": [{"keyword": w, "count": c} for w, c in pos.most_common(n)],
            "negative": [{"keyword": w, "count": c} for w, c in neg.most_common(n)],
        }
    return out


def monthly_rating_trend(apps: list[dict], reviews: list[dict]) -> dict:
    out = {}
    unparsed_total = 0
    for a in apps:
        buckets: dict[str, list[int]] = {}
        unparsed = 0
        for r in reviews:
            if r["app_id"] != a["id"]:
                continue
            created = r.get("created_at")
            if not isinstance(created, datetime):
                unparsed += 1
                continue
            month_key = created.strftime("%Y-%m")
            buckets.setdefault(month_key, []).append(r["rating"])

        series = []
        for month_key in sorted(buckets):
            ratings = buckets[month_key]
            series.append({
                "month": month_key,
                "average": round(sum(ratings) / len(ratings), 2),
                "review_count": len(ratings),
                "noise": len(ratings) < MONTHLY_NOISE_THRESHOLD,
            })
        out[a["id"]] = {"series": series, "unparsed_date_count": unparsed}
        unparsed_total += unparsed
    out["_meta"] = {"noise_threshold": MONTHLY_NOISE_THRESHOLD, "total_unparsed_dates": unparsed_total}
    return out


def sentiment_rating_mismatch(apps: list[dict], reviews: list[dict], example_limit: int = 3) -> dict:
    out = {}
    for a in apps:
        app_reviews = [r for r in reviews if r["app_id"] == a["id"]]
        high_neg = [r for r in app_reviews if r["rating"] >= 4 and r["sentiment_score"] < 0]
        low_pos = [r for r in app_reviews if r["rating"] <= 2 and r["sentiment_score"] > 0]
        out[a["id"]] = {
            "high_rating_negative_text_count": len(high_neg),
            "low_rating_positive_text_count": len(low_pos),
            "high_rating_negative_examples": [r["content"][:120] for r in high_neg[:example_limit]],
            "low_rating_positive_examples": [r["content"][:120] for r in low_pos[:example_limit]],
        }
    out["_meta"] = {
        "method": "word_dictionary_based_estimate",
        "disclaimer": "감성분석은 단어 기반 추정입니다. 정확한 감성 분류가 아닌 참고 지표로 사용하세요.",
    }
    return out


def run_full_analysis(apps: list[dict], reviews: list[dict]) -> dict:
    return {
        "rating_distribution": rating_distribution(apps, reviews),
        "wordcloud_frequencies": wordcloud_frequencies(apps, reviews),
        "top_keywords": top_keywords(apps, reviews),
        "common_unique_keywords": common_and_unique_keywords(apps, reviews),
        "positive_negative_keywords": positive_negative_keywords(apps, reviews),
        "monthly_rating_trend": monthly_rating_trend(apps, reviews),
        "sentiment_rating_mismatch": sentiment_rating_mismatch(apps, reviews),
    }


# =====================================================================================
# F5: LLM 인사이트 (Claude API)
# =====================================================================================

INSIGHT_INSTRUCTIONS = """당신은 앱 리뷰 데이터 분석가입니다. 아래는 은행/핀테크 앱 리뷰를 자동 분석한 수치 요약입니다.
각 앱에 대해 강점(strengths), 약점(weaknesses), 개선 제안(suggestions)을 한국어로 2~4개씩 bullet 형태로 작성하세요.
반드시 주어진 수치/키워드에 근거해서만 작성하고, 근거로 사용한 키워드를 evidence_keywords에 포함하세요.
아래 JSON 스키마로만 응답하세요 (설명 문장 없이 JSON만 출력):

{"apps": [{"app_id": <int>, "strengths": ["..."], "weaknesses": ["..."], "suggestions": ["..."], "evidence_keywords": ["..."]}]}
"""


def _build_insight_summary(apps: list[dict], analysis: dict) -> dict:
    rd, tk, pn, mismatch = (
        analysis["rating_distribution"], analysis["top_keywords"],
        analysis["positive_negative_keywords"], analysis["sentiment_rating_mismatch"],
    )
    summary = []
    for a in apps:
        aid = a["id"]
        summary.append({
            "app_id": aid,
            "app_name": a["name"],
            "average_rating": rd.get(aid, {}).get("average"),
            "review_count": rd.get(aid, {}).get("review_count"),
            "top_keywords": [k["keyword"] for k in tk.get(aid, [])[:10]],
            "positive_keywords": [k["keyword"] for k in pn.get(aid, {}).get("positive", [])[:5]],
            "negative_keywords": [k["keyword"] for k in pn.get(aid, {}).get("negative", [])[:5]],
            "rating_text_mismatch_count": (
                mismatch.get(aid, {}).get("high_rating_negative_text_count", 0)
                + mismatch.get(aid, {}).get("low_rating_positive_text_count", 0)
            ),
        })
    return {"apps": summary}


def generate_insights(apps: list[dict], analysis: dict, api_key: str, model: str) -> dict:
    if not api_key:
        return {
            a["id"]: {
                "strengths": [], "weaknesses": [], "suggestions": [], "evidence_keywords": [],
                "available": False,
                "note": "ANTHROPIC_API_KEY가 설정되지 않아 LLM 인사이트를 생성하지 못했습니다. (참고)",
            }
            for a in apps
        }

    summary = _build_insight_summary(apps, analysis)
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=model,
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": INSIGHT_INSTRUCTIONS + "\n\n분석 요약:\n" + json.dumps(summary, ensure_ascii=False),
            }],
        )
        raw_text = "".join(b.text for b in message.content if getattr(b, "type", None) == "text")
        parsed = json.loads(raw_text)
        out = {}
        for item in parsed.get("apps", []):
            out[item["app_id"]] = {
                "strengths": item.get("strengths", []),
                "weaknesses": item.get("weaknesses", []),
                "suggestions": item.get("suggestions", []),
                "evidence_keywords": item.get("evidence_keywords", []),
                "available": True,
                "note": "LLM이 생성한 요약입니다. 참고용으로 활용하세요.",
            }
        return out
    except Exception as e:  # noqa: BLE001 - LLM 실패가 전체 파이프라인을 막으면 안 됨
        return {
            a["id"]: {
                "strengths": [], "weaknesses": [], "suggestions": [], "evidence_keywords": [],
                "available": False,
                "note": f"LLM 인사이트 생성에 실패했습니다: {e} (참고)",
            }
            for a in apps
        }


# =====================================================================================
# F6: 내보내기 (CSV / Word / PDF)
# =====================================================================================

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager, rcParams
    if WORDCLOUD_FONT_PATH and os.path.exists(WORDCLOUD_FONT_PATH):
        rcParams["font.family"] = font_manager.FontProperties(fname=WORDCLOUD_FONT_PATH).get_name()
        rcParams["axes.unicode_minus"] = False
    _FONT_OK = True
except Exception:  # noqa: BLE001 - 폰트/matplotlib 문제가 있어도 CSV 내보내기는 동작해야 함
    _FONT_OK = False


def reviews_to_csv(rows: list[dict], processed: bool) -> bytes:
    cols = ["app_name", "rating", "content", "author", "created_at"]
    if processed:
        cols += ["clean_content", "nouns", "sentiment_score"]
    df = pd.DataFrame(rows)
    for c in cols:
        if c not in df.columns:
            df[c] = None
    if processed and "nouns" in df.columns:
        df["nouns"] = df["nouns"].apply(lambda x: " ".join(x) if isinstance(x, list) else x)
    buf = io.StringIO()
    df[cols].to_csv(buf, index=False, encoding="utf-8-sig")
    return buf.getvalue().encode("utf-8-sig")


def _fig_to_png_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def chart_rating_distribution_png(apps: list[dict], analysis: dict) -> bytes:
    rd = analysis["rating_distribution"]
    fig, ax = plt.subplots(figsize=(8, 5))
    width = 0.8 / max(len(apps), 1)
    scores = [1, 2, 3, 4, 5]
    for i, a in enumerate(apps):
        counts = rd.get(a["id"], {}).get("counts", {})
        avg = rd.get(a["id"], {}).get("average", 0)
        x = [s + (i - (len(apps) - 1) / 2) * width for s in scores]
        ax.bar(x, [counts.get(s, 0) for s in scores], width=width, label=f"{a['name']} (평균 {avg})")
    ax.set_xticks(scores)
    ax.set_xlabel("별점")
    ax.set_ylabel("리뷰 수")
    ax.set_title("앱별 별점 분포")
    ax.legend()
    return _fig_to_png_bytes(fig)


def chart_top_keywords_png(app: dict, kws: list[dict]) -> bytes:
    words = [k["keyword"] for k in kws][::-1]
    counts = [k["count"] for k in kws][::-1]
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.barh(words, counts)
    ax.set_title(f"{app['name']} 상위 키워드")
    return _fig_to_png_bytes(fig)


def chart_wordcloud_png(app: dict, freq: dict) -> bytes | None:
    if not _FONT_OK or not freq:
        return None
    try:
        from wordcloud import WordCloud
        wc = WordCloud(font_path=WORDCLOUD_FONT_PATH, background_color="white", width=600, height=400)
        wc.generate_from_frequencies(freq)
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.imshow(wc)
        ax.axis("off")
        ax.set_title(app["name"])
        return _fig_to_png_bytes(fig)
    except Exception:
        return None


def chart_monthly_trend_png(apps: list[dict], analysis: dict) -> bytes | None:
    trend = analysis["monthly_rating_trend"]
    if not any(trend.get(a["id"], {}).get("series") for a in apps):
        return None
    fig, ax = plt.subplots(figsize=(9, 5))
    for a in apps:
        series = trend.get(a["id"], {}).get("series", [])
        if series:
            ax.plot([s["month"] for s in series], [s["average"] for s in series], marker="o", label=a["name"])
    ax.set_xlabel("월")
    ax.set_ylabel("평균 별점")
    ax.set_title("월별 평균 별점 추이")
    ax.legend()
    plt.xticks(rotation=45)
    return _fig_to_png_bytes(fig)


def build_docx_report(apps: list[dict], analysis: dict, insights: dict) -> bytes:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("ReviewLens 분석 보고서", level=0)
    doc.add_paragraph(
        "이 보고서는 앱 스토어 리뷰를 자동 수집·분석한 결과입니다. "
        "감성분석은 단어 기반 추정이며, 표본이 적은 구간은 노이즈로 표시됩니다."
    )

    doc.add_heading("앱별 요약", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "앱", "평균 별점", "리뷰 수"
    for a in apps:
        rd = analysis["rating_distribution"].get(a["id"], {})
        row = table.add_row().cells
        row[0].text = a["name"]
        row[1].text = str(rd.get("average", "-"))
        row[2].text = str(rd.get("review_count", "-"))

    doc.add_heading("별점 분포", level=1)
    doc.add_picture(io.BytesIO(chart_rating_distribution_png(apps, analysis)), width=Inches(6))

    trend_png = chart_monthly_trend_png(apps, analysis)
    if trend_png:
        doc.add_heading("월별 평균 별점 추이", level=1)
        doc.add_picture(io.BytesIO(trend_png), width=Inches(6))

    cu = analysis["common_unique_keywords"]
    doc.add_heading("공통/고유 키워드", level=1)
    if cu.get("note"):
        doc.add_paragraph(cu["note"])
    else:
        doc.add_paragraph("공통 키워드: " + ", ".join(cu["common"]))
        for a in apps:
            doc.add_paragraph(f"{a['name']} 고유 키워드: " + ", ".join(cu["unique"].get(a["id"], [])))

    for a in apps:
        doc.add_heading(f"{a['name']} 상세", level=1)

        wc_png = chart_wordcloud_png(a, analysis["wordcloud_frequencies"].get(a["id"], {}))
        if wc_png:
            doc.add_picture(io.BytesIO(wc_png), width=Inches(5))

        tk = analysis["top_keywords"].get(a["id"], [])
        if tk:
            doc.add_picture(io.BytesIO(chart_top_keywords_png(a, tk)), width=Inches(5))

        mismatch = analysis["sentiment_rating_mismatch"].get(a["id"], {})
        doc.add_paragraph(
            f"감성-별점 불일치: 별점 높은데 부정 텍스트 {mismatch.get('high_rating_negative_text_count', 0)}건, "
            f"별점 낮은데 긍정 텍스트 {mismatch.get('low_rating_positive_text_count', 0)}건 (단어 기반 추정, 참고용)"
        )

        ins = insights.get(a["id"], {})
        doc.add_heading("LLM 인사이트 (참고)", level=2)
        if not ins.get("available"):
            doc.add_paragraph(ins.get("note", "인사이트를 사용할 수 없습니다."))
        else:
            for label, key in [("강점", "strengths"), ("약점", "weaknesses"), ("개선 제안", "suggestions")]:
                doc.add_paragraph(label + ":", style="Intense Quote")
                for line in ins.get(key, []):
                    doc.add_paragraph(line, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf_report(apps: list[dict], analysis: dict, insights: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table
    from reportlab.lib.styles import getSampleStyleSheet

    styles = getSampleStyleSheet()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    story = [Paragraph("ReviewLens 분석 보고서", styles["Title"]), Spacer(1, 12)]
    story.append(Paragraph("감성분석은 단어 기반 추정이며, 표본이 적은 구간은 노이즈로 표시됩니다.", styles["Normal"]))
    story.append(Spacer(1, 12))

    data = [["앱", "평균 별점", "리뷰 수"]]
    for a in apps:
        rd = analysis["rating_distribution"].get(a["id"], {})
        data.append([a["name"], str(rd.get("average", "-")), str(rd.get("review_count", "-"))])
    story.append(Table(data))
    story.append(Spacer(1, 12))

    story.append(Image(io.BytesIO(chart_rating_distribution_png(apps, analysis)), width=16 * cm, height=10 * cm))
    story.append(Spacer(1, 12))

    trend_png = chart_monthly_trend_png(apps, analysis)
    if trend_png:
        story.append(Image(io.BytesIO(trend_png), width=16 * cm, height=9 * cm))
        story.append(Spacer(1, 12))

    for a in apps:
        story.append(Paragraph(f"{a['name']} — LLM 인사이트 (참고)", styles["Heading2"]))
        ins = insights.get(a["id"], {})
        if not ins.get("available"):
            story.append(Paragraph(ins.get("note", "인사이트를 사용할 수 없습니다."), styles["Normal"]))
        else:
            for label, key in [("강점", "strengths"), ("약점", "weaknesses"), ("개선 제안", "suggestions")]:
                story.append(Paragraph(f"<b>{label}</b>: " + " / ".join(ins.get(key, [])), styles["Normal"]))
        story.append(Spacer(1, 10))

    doc.build(story)
    return buf.getvalue()


# =====================================================================================
# 파이프라인 실행 (F2~F5를 순서대로, 앱 수집만 병렬)
# =====================================================================================

def run_pipeline(raw_inputs: list[str], count: int, api_key: str, model: str, progress_box) -> dict:
    apps: list[dict] = []
    skipped_duplicates: list[str] = []
    seen_pkg: dict[str, int] = {}
    for raw in raw_inputs:
        try:
            pkg = parse_package_id(raw)
        except ValueError as e:
            return {"error": str(e)}
        if pkg in seen_pkg:
            skipped_duplicates.append(raw)
            continue
        info = fetch_app_info(pkg)
        aid = len(apps) + 1
        apps.append({"id": aid, "package_id": pkg, "name": info["name"], "icon_url": info["icon_url"]})
        seen_pkg[pkg] = aid

    if skipped_duplicates:
        progress_box.info(f"중복된 앱 {len(skipped_duplicates)}개는 건너뛰었습니다: {', '.join(skipped_duplicates)}")

    bars = {a["id"]: progress_box.progress(0, text=f"{a['name']}: 대기 중") for a in apps}

    collected_rows: list[dict] = []
    succeeded_ids: list[int] = []
    failed: dict[int, str] = {}

    with concurrent.futures.ThreadPoolExecutor(max_workers=max(1, len(apps))) as executor:
        future_to_app = {executor.submit(fetch_reviews_with_retry, a["package_id"], count): a for a in apps}
        for future in concurrent.futures.as_completed(future_to_app):
            a = future_to_app[future]
            try:
                rows = future.result()
                for row in rows:
                    row["app_id"] = a["id"]
                collected_rows.extend(rows)
                succeeded_ids.append(a["id"])
                bars[a["id"]].progress(100, text=f"✅ {a['name']}: {len(rows)}건 수집 완료")
            except Exception as e:
                failed[a["id"]] = str(e)
                bars[a["id"]].progress(100, text=f"❌ {a['name']}: 수집 실패")

    if not succeeded_ids:
        return {"error": "모든 앱의 리뷰 수집에 실패했습니다.", "failed": failed}

    rows_for_processing = [r for r in collected_rows if r["app_id"] in succeeded_ids]
    processed = preprocess_reviews(rows_for_processing)
    for row in processed:
        row["sentiment_score"] = compute_sentiment(row["content"])

    if not processed:
        return {"error": "전처리 후 남은 리뷰가 없습니다 (모두 중복·초단문·명사 없음으로 제거됨).", "failed": failed}

    apps_with_data = {r["app_id"] for r in processed}
    for aid in succeeded_ids:
        if aid not in apps_with_data:
            failed[aid] = "전처리 후 분석 가능한 리뷰가 0건입니다 (중복·초단문·특수문자뿐인 리뷰만 수집됨)."
    succeeded_ids = [aid for aid in succeeded_ids if aid in apps_with_data]

    if not succeeded_ids:
        return {"error": "모든 앱이 전처리 후 분석 가능한 리뷰가 남지 않았습니다.", "failed": failed}

    succeeded_apps = [a for a in apps if a["id"] in succeeded_ids]

    analysis_result = run_full_analysis(succeeded_apps, processed)
    insights_result = generate_insights(succeeded_apps, analysis_result, api_key, model)

    return {
        "apps": succeeded_apps,
        "failed": failed,
        "raw_rows": collected_rows,
        "processed_rows": processed,
        "analysis": analysis_result,
        "insights": insights_result,
    }


# =====================================================================================
# Streamlit UI
# =====================================================================================

st.set_page_config(page_title="ReviewLens", page_icon="📱", layout="wide")

if "result" not in st.session_state:
    st.session_state.result = None

st.title("📱 ReviewLens")
st.caption("앱 스토어 리뷰를 자동 수집·분석해 경쟁 비교와 개선 인사이트를 제공합니다.")

with st.sidebar:
    st.header("분석할 앱")
    apps_text = st.text_area(
        f"스토어 URL 또는 패키지명 (줄바꿈으로 최대 {MAX_APPS}개)",
        height=130,
        placeholder="https://play.google.com/store/apps/details?id=viva.republica.toss\nkr.co.busanbank.mbp",
    )
    count = st.number_input("앱당 수집 개수", min_value=10, max_value=1000, value=DEFAULT_COUNT, step=10)

    st.divider()
    st.caption("F5 LLM 인사이트 (선택 — 비워두면 인사이트 없이 나머지 분석만 진행)")
    api_key_input = st.text_input("ANTHROPIC_API_KEY", type="password", value=DEFAULT_ANTHROPIC_KEY)
    model_input = st.text_input("모델", value=DEFAULT_ANTHROPIC_MODEL)

    st.divider()
    start = st.button("분석 시작", type="primary", use_container_width=True)

if start:
    urls = [u.strip() for u in apps_text.splitlines() if u.strip()]
    if not urls:
        st.error("최소 1개 이상의 앱을 입력하세요.")
    elif len(urls) > MAX_APPS:
        st.error(f"한 번에 최대 {MAX_APPS}개 앱까지 분석할 수 있습니다.")
    else:
        progress_box = st.container()
        with st.spinner("리뷰 수집 → 전처리 → 분석 → 인사이트 생성 중..."):
            st.session_state.result = run_pipeline(urls, int(count), api_key_input, model_input, progress_box)

result = st.session_state.result

if result is None:
    st.info("왼쪽 사이드바에서 앱을 입력하고 '분석 시작'을 눌러주세요.")
elif result.get("error"):
    st.error(result["error"])
    if result.get("failed"):
        for aid, msg in result["failed"].items():
            st.caption(f"- 앱 #{aid}: {msg}")
else:
    apps = result["apps"]
    analysis = result["analysis"]
    insights = result["insights"]

    if result.get("failed"):
        st.warning("일부 앱은 실패해 결과에서 제외됐습니다: " + ", ".join(
            f"#{aid} ({msg})" for aid, msg in result["failed"].items()
        ))

    cols = st.columns(len(apps))
    for col, a in zip(cols, apps):
        rd = analysis["rating_distribution"].get(a["id"], {})
        col.metric(a["name"], rd.get("average", "-"), f"{rd.get('review_count', 0)}건 리뷰")

    tabs = st.tabs([
        "별점 분포", "워드클라우드", "상위 키워드", "공통/고유 키워드",
        "긍정/부정 키워드", "월별 추이", "감성-별점 불일치",
    ])

    with tabs[0]:
        rd = analysis["rating_distribution"]
        df = pd.DataFrame(
            {a["name"]: [rd[a["id"]]["counts"][s] for s in range(1, 6)] for a in apps},
            index=[f"{s}점" for s in range(1, 6)],
        )
        st.bar_chart(df)

    with tabs[1]:
        wc = analysis["wordcloud_frequencies"]
        for a in apps:
            st.markdown(f"**{a['name']}**")
            freq = wc.get(a["id"], {})
            items = sorted(freq.items(), key=lambda x: -x[1])[:60]
            if not items:
                st.caption("표시할 키워드가 없습니다.")
                continue
            max_c, min_c = items[0][1], items[-1][1]
            spans = []
            for word, c in items:
                scale = 0.9 if max_c == min_c else 0.9 + (c - min_c) / (max_c - min_c) * 1.6
                spans.append(
                    f'<span title="{word}: {c}회" '
                    f'style="font-size:{scale:.2f}rem;margin:4px;display:inline-block;">{word}</span>'
                )
            st.markdown(" ".join(spans), unsafe_allow_html=True)

    with tabs[2]:
        tk = analysis["top_keywords"]
        for a in apps:
            st.markdown(f"**{a['name']}**")
            kws = tk.get(a["id"], [])
            if kws:
                df = pd.DataFrame(kws).set_index("keyword")
                st.bar_chart(df)
            else:
                st.caption("표시할 키워드가 없습니다.")

    with tabs[3]:
        cu = analysis["common_unique_keywords"]
        if cu.get("note"):
            st.info(cu["note"])
        else:
            st.markdown("**공통 키워드**: " + ", ".join(cu["common"]))
            for a in apps:
                st.markdown(f"**{a['name']} 고유 키워드**: " + ", ".join(cu["unique"].get(a["id"], [])))

    with tabs[4]:
        pn = analysis["positive_negative_keywords"]
        for a in apps:
            st.markdown(f"**{a['name']}**")
            c1, c2 = st.columns(2)
            pos = pn.get(a["id"], {}).get("positive", [])
            neg = pn.get(a["id"], {}).get("negative", [])
            with c1:
                st.caption("긍정 리뷰(별점 4~5) 키워드")
                if pos:
                    st.bar_chart(pd.DataFrame(pos).set_index("keyword"))
            with c2:
                st.caption("부정 리뷰(별점 1~2) 키워드")
                if neg:
                    st.bar_chart(pd.DataFrame(neg).set_index("keyword"))

    with tabs[5]:
        trend = analysis["monthly_rating_trend"]
        months = sorted({p["month"] for a in apps for p in trend.get(a["id"], {}).get("series", [])})
        if not months:
            st.info("월별 추이를 계산할 만한 작성일 데이터가 없습니다.")
        else:
            df = pd.DataFrame(index=months)
            for a in apps:
                series = {p["month"]: p["average"] for p in trend.get(a["id"], {}).get("series", [])}
                df[a["name"]] = [series.get(m) for m in months]
            st.line_chart(df)
            if any(p["noise"] for a in apps for p in trend.get(a["id"], {}).get("series", [])):
                st.warning("일부 구간은 해당 월 리뷰 표본이 적어 평균 별점의 신뢰도가 낮을 수 있습니다 (노이즈 구간).")

    with tabs[6]:
        meta = analysis["sentiment_rating_mismatch"].get("_meta", {})
        st.info(meta.get("disclaimer", ""))
        for a in apps:
            m = analysis["sentiment_rating_mismatch"].get(a["id"], {})
            st.markdown(
                f"**{a['name']}** — 별점 높은데 부정 텍스트 {m.get('high_rating_negative_text_count', 0)}건, "
                f"별점 낮은데 긍정 텍스트 {m.get('low_rating_positive_text_count', 0)}건"
            )
            for t in m.get("high_rating_negative_examples", []):
                st.caption(f"★높음·부정: {t}")
            for t in m.get("low_rating_positive_examples", []):
                st.caption(f"★낮음·긍정: {t}")

    st.header("LLM 인사이트")
    st.caption("감성분석은 단어 기반 추정이며, 아래 인사이트는 LLM이 생성한 요약으로 참고용입니다.")
    for a in apps:
        ins = insights.get(a["id"], {})
        with st.expander(f"{a['name']} — 참고", expanded=True):
            if not ins.get("available"):
                st.caption(ins.get("note", "인사이트를 사용할 수 없습니다."))
            else:
                st.markdown("**강점**")
                for s in ins.get("strengths", []):
                    st.markdown(f"- {s}")
                st.markdown("**약점**")
                for s in ins.get("weaknesses", []):
                    st.markdown(f"- {s}")
                st.markdown("**개선 제안**")
                for s in ins.get("suggestions", []):
                    st.markdown(f"- {s}")
                if ins.get("evidence_keywords"):
                    st.caption("근거 키워드: " + ", ".join(ins["evidence_keywords"]))

    st.header("리포트 / 내보내기")
    name_map = {a["id"]: a["name"] for a in apps}

    raw_rows = [{**r, "app_name": name_map.get(r["app_id"], r["app_id"])} for r in result["raw_rows"]]
    processed_rows = [{**r, "app_name": name_map.get(r["app_id"], r["app_id"])} for r in result["processed_rows"]]

    c1, c2, c3, c4 = st.columns(4)
    c1.download_button(
        "원본 리뷰 CSV", data=reviews_to_csv(raw_rows, processed=False),
        file_name="reviewlens_raw.csv", mime="text/csv", use_container_width=True,
    )
    c2.download_button(
        "전처리 리뷰 CSV", data=reviews_to_csv(processed_rows, processed=True),
        file_name="reviewlens_processed.csv", mime="text/csv", use_container_width=True,
    )
    if c3.button("Word 보고서 생성", use_container_width=True):
        st.session_state["docx_bytes"] = build_docx_report(apps, analysis, insights)
    if c4.button("PDF 보고서 생성", use_container_width=True):
        st.session_state["pdf_bytes"] = build_pdf_report(apps, analysis, insights)

    if st.session_state.get("docx_bytes"):
        st.download_button(
            "⬇ Word 보고서 다운로드", data=st.session_state["docx_bytes"],
            file_name="reviewlens_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if st.session_state.get("pdf_bytes"):
        pdf_bytes = st.session_state["pdf_bytes"]
        st.download_button("⬇ PDF 보고서 다운로드", data=pdf_bytes, file_name="reviewlens_report.pdf", mime="application/pdf")
        with st.expander("PDF 미리보기"):
            b64 = base64.b64encode(pdf_bytes).decode()
            st.components.v1.html(
                f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="600"></iframe>',
                height=620,
            )
